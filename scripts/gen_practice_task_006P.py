# -*- coding: utf-8 -*-
"""Generate Practice_Task.pdf for 006_P Safety module."""

import sys
import os

# Output paths
LESSON_DIR = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "Grade_9", "Semester_1", "01_Safety",
    "006_P - Safety checklist rehearsal + common mistake review"
)
DOCX_PATH = os.path.join(LESSON_DIR, "Practice_Task.docx")
PDF_PATH = os.path.join(LESSON_DIR, "Practice_Task.pdf")

from docx import Document
from docx.shared import Pt, Cm, Emu, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── Colors ──────────────────────────────────────────────────────
NAVY = RGBColor(0x1F, 0x4E, 0x79)
BLUE = RGBColor(0x2E, 0x75, 0xB6)
GREY = RGBColor(0x80, 0x80, 0x80)
BODY_COLOR = RGBColor(0x33, 0x33, 0x33)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BORDER_COLOR = "BFBFBF"

doc = Document()

# ── Page setup: A4, 1" margins ──────────────────────────────────
for section in doc.sections:
    section.page_width = Emu(11906 * 914)  # A4 width
    section.page_height = Emu(16838 * 914)  # A4 height
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# ── Default font ────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)
font.color.rgb = BODY_COLOR
pf = style.paragraph_format
pf.space_after = Pt(80 / 20)  # 80 TWIPs = 4pt
pf.line_spacing = 1.15


def add_para(text, size=11, color=BODY_COLOR, bold=False, alignment=None,
             space_before=None, space_after=None, font_name='Arial'):
    """Add a paragraph with specified formatting."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    if alignment is not None:
        p.alignment = alignment
    pf = p.paragraph_format
    if space_before is not None:
        pf.space_before = Twips(space_before)
    if space_after is not None:
        pf.space_after = Twips(space_after)
    pf.line_spacing = 1.15
    return p


def add_h2(text):
    """Section heading (H2)."""
    return add_para(text, size=13, color=NAVY, bold=True,
                    space_before=300, space_after=120)


def add_h3(text):
    """Topic group heading (H3)."""
    return add_para(text, size=11.5, color=BLUE, bold=True,
                    space_before=240, space_after=80)


def add_body(text, space_after_twips=80):
    """Body paragraph."""
    return add_para(text, size=11, color=BODY_COLOR,
                    space_after=space_after_twips)


def add_mcq_option(letter, text, space_after_twips=40):
    """MCQ option line, indented."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Twips(360)
    pf.space_after = Twips(space_after_twips)
    pf.line_spacing = 1.15
    run = p.add_run(f"{letter}. {text}")
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.font.color.rgb = BODY_COLOR
    return p


def add_checklist_item(text):
    """Checklist item with ☐."""
    p = add_para(f"\u2610 {text}", size=11, color=BODY_COLOR, space_after=60)
    return p


def add_horizontal_rule():
    """Navy horizontal rule."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Twips(60)
    pf.space_after = Twips(80)
    # Bottom border on the paragraph
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="8" w:space="1" w:color="1F4E79"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    return p


def make_revision_table(rows):
    """Create the revision pointers table."""
    table = doc.add_table(rows=1 + len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Twips(3000)
        row.cells[1].width = Twips(6026)

    # Header row
    hdr = table.rows[0]
    for i, text in enumerate(["Tema", "Kur peržiūrėti"]):
        cell = hdr.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        run.font.color.rgb = WHITE
        run.bold = True
        # Navy background
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="1F4E79" w:val="clear"/>'
        )
        cell._tc.get_or_add_tcPr().append(shading)

    # Data rows
    for r_idx, (tema, kur) in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, text in enumerate([tema, kur]):
            cell = row.cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.font.name = 'Arial'
            run.font.size = Pt(11)
            run.font.color.rgb = BODY_COLOR

    # Borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(
        f'<w:tblPr {nsdecls("w")}/>').append(tbl)
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{BORDER_COLOR}"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="{BORDER_COLOR}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{BORDER_COLOR}"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="{BORDER_COLOR}"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{BORDER_COLOR}"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="{BORDER_COLOR}"/>'
        f'</w:tblBorders>'
    )
    tblPr = tbl.tblPr
    # Remove existing borders if any
    for child in list(tblPr):
        if child.tag.endswith('tblBorders'):
            tblPr.remove(child)
    tblPr.append(borders)

    return table


# ═══════════════════════════════════════════════════════════════
# 1. HEADER
# ═══════════════════════════════════════════════════════════════

# Label line
add_para("PRAKTINĖS UŽDUOTYS", size=9, color=GREY,
         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=60)

# Title
add_para("Saugos praktikos užduotys", size=18, color=NAVY, bold=True,
         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=60)

# Metadata line
add_para("9 klasė  \u2022  Sauga  \u2022  tipas \u201EP\u201C", size=10, color=GREY,
         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=80)

# Horizontal rule
add_horizontal_rule()

# ═══════════════════════════════════════════════════════════════
# 2. KĄ PADARYSITE
# ═══════════════════════════════════════════════════════════════

add_h2("Ką padarysite")

add_body(
    "Šiomis užduotimis pasiruošite saugos modulio atsiskaitomajam darbui (007_A). "
    "Pasitikrinsite, ką prisimenate iš ergonomikos, privatumo, internetinių grėsmių "
    "ir aplinkos poveikio pamokų (001_L, 002_L, 003_L, 004_L). "
    "Jei kai kurių atsakymų neprisimenat, skyriuje \u201EKą daryti, jei sunku\u201C "
    "rasite nuorodas į konkrečias temas."
)

# ═══════════════════════════════════════════════════════════════
# 3. UŽDUOTYS
# ═══════════════════════════════════════════════════════════════

add_h2("Užduotys")

# ── Topic 1: Ergonomika ────────────────────────────────────────
add_h3("1. Ergonomika")

# Q1 — MCQ (T2+T4: context shift home + added complexity)
add_body(
    "1. Jūsų klasiokas ruošia namų darbus ant lovos su nešiojamu kompiuteriu. "
    "Jis sėdi pasilenkęs, sukryžiavęs kojas. "
    "Kuris teiginys geriausiai paaiškina, kodėl tokia darbo padėtis yra kenksminga?"
)
add_mcq_option("A", "Ant minkšto paviršiaus kompiuteris greičiau perkaista ir pradeda veikti lėčiau")
add_mcq_option("B", "Tokioje padėtyje rankos yra per arti klaviatūros ir pirštai greičiau pavargsta")
add_mcq_option("C", "Ilgai sėdint netaisyklingai apkraunami nugaros ir kaklo raumenys, atsiranda skausmas")
add_mcq_option("D", "Lovoje šviesa yra per silpna, todėl akys labiau įsitempia ir greičiau pavargsta", space_after_twips=80)

# Q2 — Composite essay (T2+T3: 20-20-20 + ergonomics evaluation merged)
add_body(
    "2. Jūsų draugas sako: \u201E20-20-20 taisyklė yra nesąmonė, nes aš per pertrauką "
    "vis tiek žiūriu į telefoną.\u201C "
    "Paaiškinkite, kodėl draugas klysta. Nurodykite, kas turi vykti per tą 20 sekundžių "
    "poilsį ir kodėl žiūrėjimas į telefoną to nepakeičia. Atsakykite 2\u20133 sakiniais."
)

# ── Topic 2: Privatumas ir paskyrų sauga ──────────────────────
add_h3("2. Privatumas ir paskyrų sauga")

# Q3 — MCQ (T2: closer distractors, partially correct option)
add_body(
    "3. Kuris iš šių slaptažodžių yra saugiausias?"
)
add_mcq_option("A", "mK8#zLp2!qXw")
add_mcq_option("B", "Vilnius2024!")
add_mcq_option("C", "ManoKatinasRudas")
add_mcq_option("D", "qwerty12345abc", space_after_twips=80)

# Q4 — Composite scenario (T3: 2FA + personal data merged)
add_body(
    "4. Jūsų draugas registruojasi naujoje svetainėje. Jis naudoja stiprų slaptažodį, "
    "bet atsisako įjungti dviejų veiksnių autentifikaciją (2FA), nes \u201Etai per "
    "sudėtinga.\u201C Registracijos formoje jis nurodo tikrą vardą, pavardę, gimimo datą "
    "ir mokyklos pavadinimą. Atlikite du dalykus: "
    "(a) paaiškinkite, kodėl 2FA yra svarbi net turint stiprų slaptažodį, ir "
    "(b) nurodykite, kuriuos iš šių duomenų buvo saugu nurodyti, o kuriuos ne. "
    "Atsakykite 3\u20134 sakiniais."
)

# ── Topic 3: Internetinės grėsmės ─────────────────────────────
add_h3("3. Internetinės grėsmės")

# Q5 — MCQ (T2+T4: context shift to social media + closer distractors)
add_body(
    "5. Socialiniame tinkle gaunate žinutę nuo nepažįstamo žmogaus: "
    "\u201ESveiki! Radau jūsų nuotrauką keistoje svetainėje. "
    "Pažiūrėkite čia: [nuoroda].\u201C "
    "Kuris šios žinutės požymis yra patikimiausias phishing rodiklis?"
)
add_mcq_option("A", "Žinutė parašyta be rašybos klaidų ir atrodo tvarkingai")
add_mcq_option("B", "Siuntėjas yra nepažįstamas ir ragina paspausti nuorodą")
add_mcq_option("C", "Žinutė atsiųsta vakare, ne įprastu bendravimo metu")
add_mcq_option("D", "Žinutės tekste nėra jūsų tikrojo vardo, tik bendras kreipinys", space_after_twips=80)

# Q6 — MCQ classify+justify (T1: Bloom's +1 Understand→Apply)
add_body(
    "6. Perskaitykite situaciją ir pasirinkite teisingą atsakymą. "
    "Klasiokė socialiniame tinkle pasidalino straipsniu, kuriame teigiama, "
    "kad telefonai kenkia sveikatai. Ji tiki, kad tai tiesa, nors mokslininkai "
    "to nepatvirtino. Kaip teisingai klasifikuoti šį atvejį?"
)
add_mcq_option("A", "Tai yra dezinformacija, nes straipsnis yra neteisingas ir klasiokė jį platina toliau")
add_mcq_option("B", "Tai yra kibernetinė ataka, nes straipsnis siekia paveikti žmonių sveikatą ir elgesį")
add_mcq_option("C", "Tai nėra nei dezinformacija, nei klaidinga informacija, nes klasiokė nieko blogo nedarė")
add_mcq_option("D", "Tai yra klaidinga informacija, nes klasiokė platina netikslų turinį netyčia, nesiekdama apgauti", space_after_twips=80)

# Q7 — Composite essay (T3: phishing recognition + safe response merged)
add_body(
    "7. Per el. paštą gaunate laišką nuo \u201Emokyklos administracijos\u201C "
    "(adresas: mokykla-info@gmail.com). Laiške rašoma: \u201EPrašome per 24 valandas "
    "atnaujinti savo prisijungimo duomenis paspaudę nuorodą. Priešingu atveju paskyra "
    "bus užblokuota.\u201C "
    "Atlikite du dalykus: (a) nurodykite bent du požymius, kodėl šis laiškas gali būti "
    "phishing ataka, ir (b) aprašykite, kaip elgtumėtės pagal trijų žingsnių algoritmą "
    "(sustoti, patikrinti, pranešti). Atsakykite 4\u20135 sakiniais."
)

# ── Topic 4: Skaitmeninių technologijų poveikis aplinkai ──────
add_h3("4. Skaitmeninių technologijų poveikis aplinkai")

# Q8 — MCQ (T4: context shift to personal daily actions)
add_body(
    "8. Kuris kasdieninis veiksmas sunaudoja daugiausiai energijos duomenų centruose?"
)
add_mcq_option("A", "Valandos trukmės vaizdo įrašo žiūrėjimas aukšta raiška internete")
add_mcq_option("B", "Dešimties el. laiškų be priedų išsiuntimas per vieną dieną")
add_mcq_option("C", "Pusvalandžio naršymas tekstiniame tinklalapyje be paveikslėlių")
add_mcq_option("D", "Vienos nuotraukos įkėlimas į socialinį tinklą su trumpu aprašymu", space_after_twips=80)

# Q9 — Essay (T1: Understand→Apply, propose concrete actions)
add_body(
    "9. Pamokoje mokėtės, kas yra skaitmeninis pėdsakas. "
    "Įsivaizduokite, kad norite jį sumažinti. "
    "Pasiūlykite tris konkrečius kasdienius veiksmus ir trumpai paaiškinkite, "
    "kodėl kiekvienas mažina poveikį aplinkai. Atsakykite sunumeruotu sąrašu."
)

# Q10 — Composite essay (T3: e-waste + positive ST use merged)
add_body(
    "10. Paaiškinkite, kodėl netinkamai išmesti elektroniniai įrenginiai (e-atliekos) "
    "kenkia aplinkai. Tada pateikite vieną pavyzdį, kaip skaitmeninės technologijos, "
    "priešingai, padeda spręsti aplinkosaugos problemas. "
    "Atsakykite 3\u20134 sakiniais."
)

# ═══════════════════════════════════════════════════════════════
# 4. PATIKRINKITE SAVE
# ═══════════════════════════════════════════════════════════════

add_h2("Patikrinkite save")

add_checklist_item("Ar galiu užtikrintai paaiškinti, kaip taisyklingai įrengti darbo vietą prie kompiuterio?")
add_checklist_item("Ar galiu užtikrintai apibūdinti 20-20-20 taisyklę ir paaiškinti, kodėl ji veikia?")
add_checklist_item("Ar galiu užtikrintai nurodyti bent tris stipraus slaptažodžio požymius?")
add_checklist_item("Ar galiu užtikrintai paaiškinti, kaip 2FA apsaugo paskyrą net pavogus slaptažodį?")
add_checklist_item("Ar galiu užtikrintai atpažinti phishing žinutės požymius?")
add_checklist_item("Ar galiu užtikrintai atskirti dezinformaciją nuo klaidingos informacijos?")
add_checklist_item("Ar galiu užtikrintai pritaikyti saugaus reagavimo algoritmą (sustoti, patikrinti, pranešti)?")
add_checklist_item("Ar galiu užtikrintai paaiškinti, kaip kasdieniai skaitmeniniai veiksmai veikia aplinką?")

# ═══════════════════════════════════════════════════════════════
# 5. KĄ DARYTI, JEI SUNKU
# ═══════════════════════════════════════════════════════════════

add_h2("Ką daryti, jei sunku")

make_revision_table([
    ("Ergonomika",
     "001_L, Theory_Pack: sėdėjimo laikysena, ekrano padėtis, 20-20-20 taisyklė"),
    ("Privatumas ir paskyrų sauga",
     "002_L, Theory_Pack: stiprus slaptažodis, 2FA, asmeniniai duomenys"),
    ("Internetinės grėsmės",
     "003_L, Theory_Pack: phishing požymiai, dezinformacija, sustoti-patikrinti-pranešti"),
    ("Skaitmeninių technologijų poveikis aplinkai",
     "004_L, Theory_Pack: energijos suvartojimas, e-atliekos, skaitmeninis pėdsakas"),
])

# ═══════════════════════════════════════════════════════════════
# Save and convert
# ═══════════════════════════════════════════════════════════════

doc.save(DOCX_PATH)
print(f"DOCX saved: {DOCX_PATH}")

from docx2pdf import convert
convert(DOCX_PATH, PDF_PATH)

if os.path.exists(PDF_PATH) and os.path.getsize(PDF_PATH) > 0:
    print(f"PDF saved: {PDF_PATH} ({os.path.getsize(PDF_PATH)} bytes)")
    os.remove(DOCX_PATH)
    print("Intermediate DOCX deleted.")
else:
    print("ERROR: PDF generation failed or file is empty.")
    sys.exit(1)
