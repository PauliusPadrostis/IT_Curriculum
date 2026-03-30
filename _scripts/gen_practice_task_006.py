"""
Generate Practice_Task.docx for 006_P — Safety checklist rehearsal.
Uses python-docx. Output: Practice_Task.docx → convert to PDF via docx2pdf.
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Twips, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

OUTPUT_DIR = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "Grade_9", "Semester_1", "01_Safety",
    "006_P - Safety checklist rehearsal + common mistake review",
)

NAVY = RGBColor(0x1F, 0x4E, 0x79)
BLUE = RGBColor(0x2E, 0x75, 0xB6)
GREY = RGBColor(0x80, 0x80, 0x80)
BODY_COLOR = RGBColor(0x33, 0x33, 0x33)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BORDER_GREY = "BFBFBF"

# ── helpers ──────────────────────────────────────────────────────

def set_font(run, name="Arial", size=11, color=BODY_COLOR, bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold


def add_paragraph(doc, text, font_size=11, color=BODY_COLOR, bold=False,
                  alignment=None, space_before=None, space_after=80,
                  first_line_indent=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=font_size, color=color, bold=bold)
    pf = p.paragraph_format
    if alignment is not None:
        pf.alignment = alignment
    if space_before is not None:
        pf.space_before = Twips(space_before)
    if space_after is not None:
        pf.space_after = Twips(space_after)
    pf.line_spacing = 1.15
    if first_line_indent is not None:
        pf.first_line_indent = Twips(first_line_indent)
    return p


def add_h2(doc, text):
    return add_paragraph(doc, text, font_size=13, color=NAVY, bold=True,
                         space_before=300, space_after=120)


def add_h3(doc, text):
    return add_paragraph(doc, text, font_size=11.5, color=BLUE, bold=True,
                         space_before=240, space_after=80)


def add_body(doc, text, space_after=80):
    return add_paragraph(doc, text, font_size=11, color=BODY_COLOR,
                         space_after=space_after)


def add_mcq_option(doc, letter, text, indent=360):
    p = doc.add_paragraph()
    run = p.add_run(f"{letter}. {text}")
    set_font(run, size=11, color=BODY_COLOR)
    p.paragraph_format.space_after = Twips(40)
    p.paragraph_format.left_indent = Twips(indent)
    p.paragraph_format.line_spacing = 1.15
    return p


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="1F4E79"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    p.paragraph_format.space_after = Twips(80)
    return p


def add_checklist_item(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(f"\u2610 {text}")
    set_font(run, size=11, color=BODY_COLOR)
    p.paragraph_format.space_after = Twips(60)
    p.paragraph_format.line_spacing = 1.15
    return p


def set_cell_shading(cell, color_hex):
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._element.get_or_add_tcPr().append(shading)


def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="{BORDER_GREY}"/>'
        f'<w:left w:val="single" w:sz="4" w:space="0" w:color="{BORDER_GREY}"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="{BORDER_GREY}"/>'
        f'<w:right w:val="single" w:sz="4" w:space="0" w:color="{BORDER_GREY}"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="{BORDER_GREY}"/>'
        f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="{BORDER_GREY}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def write_cell(cell, text, bold=False, color=BODY_COLOR, size=11, alignment=None):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    set_font(run, size=size, color=color, bold=bold)
    if alignment is not None:
        p.paragraph_format.alignment = alignment
    p.paragraph_format.space_after = Twips(40)
    p.paragraph_format.line_spacing = 1.15


# ── main ─────────────────────────────────────────────────────────

def build():
    doc = Document()

    # Page setup: A4, 1" margins
    section = doc.sections[0]
    section.page_width = Twips(11906)
    section.page_height = Twips(16838)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # Remove default empty paragraph
    if doc.paragraphs:
        doc.paragraphs[0]._element.getparent().remove(doc.paragraphs[0]._element)

    # ── HEADER ───────────────────────────────────────────────────
    add_paragraph(doc, "PRAKTINĖS UŽDUOTYS", font_size=9, color=GREY,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=60,
                  bold=False)

    add_paragraph(doc, "Saugos modulio kartojimas prieš vertinimą",
                  font_size=18, color=NAVY, bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=60)

    add_paragraph(doc, "9 klasė  \u2022  Sauga  \u2022  tipas \u201EP\u201C",
                  font_size=10, color=GREY,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=80)

    add_horizontal_rule(doc)

    # ── KĄ PADARYSITE ────────────────────────────────────────────
    add_h2(doc, "Ką padarysite")

    add_body(doc,
        "Šiomis užduotimis pasiruošite saugos modulio atsiskaitomajam darbui "
        "(007_A). Pasitikrinsite, ką prisimenate iš ergonomikos, privatumo, "
        "internetinių grėsmių ir aplinkos poveikio pamokų "
        "(001_L, 002_L, 003_L, 004_L)."
    )
    add_body(doc,
        "Jei kai kurių atsakymų neprisimenat, skyriuje "
        "\u201EKą daryti, jei sunku\u201C rasite nuorodas į konkrečias temas."
    )

    # ── UŽDUOTYS ─────────────────────────────────────────────────
    add_h2(doc, "Užduotys")

    # ── 1. Ergonomika ────────────────────────────────────────────
    add_h3(doc, "1. Ergonomika")

    # Q1: Technique 4 (Context Shift) — MCQ
    add_body(doc,
        "1. Jūsų šeimos narys kasdien dirba namuose su nešiojamu kompiuteriu, "
        "padėtu ant sofos. Kuris teiginys geriausiai apibūdina pagrindinę šio "
        "darbo būdo problemą?"
    )
    add_mcq_option(doc, "A",
        "Nešiojamas kompiuteris greičiau perkaista ant minkšto paviršiaus")
    add_mcq_option(doc, "B",
        "Ant sofos neįmanoma išlaikyti taisyklingos nugaros, kaklo ir rankų padėties")
    add_mcq_option(doc, "C",
        "Minkštas paviršius sumažina ekrano ryškumą ir akys greičiau pavargsta")
    add_mcq_option(doc, "D",
        "Nešiojamo kompiuterio klaviatūra ant sofos yra per toli nuo ekrano")

    # Q2: Technique 1 (Bloom's +1, Understand→Apply) — Short answer
    add_body(doc,
        "2. Jūsų klasiokas sėdi prie kompiuterio jau trečią valandą be pertraukos. "
        "Jis sėdi ant žemos kėdės, todėl jo žvilgsnis nukreiptas aukštyn į ekraną, "
        "o rankos yra aukščiau alkūnių lygio. Nurodykite visas ergonomikos klaidas "
        "šioje situacijoje ir kiekvienai pasiūlykite konkretų pataisymą. "
        "Atsakykite sunumeruotu sąrašu.",
        space_after=120
    )

    # ── 2. Privatumas ir paskyrų sauga ───────────────────────────
    add_h3(doc, "2. Privatumas ir paskyrų sauga")

    # Q3: Technique 2 (Added Complexity) — MCQ
    add_body(doc,
        "3. Kuris iš šių slaptažodžių yra saugiausias?"
    )
    add_mcq_option(doc, "A", "Vilnius2024!")
    add_mcq_option(doc, "B", "kL9#mzQp2!xR")
    add_mcq_option(doc, "C", "ManoSunis!Puskas7")
    add_mcq_option(doc, "D", "qwertyuiop!2025")

    # Q4: Technique 3 (Composite: 2FA + slaptažodžio stiprumas) — Short answer
    add_body(doc,
        "4. Jūsų draugas sako: \u201ETuriu ilgą slaptažodį, tai man 2FA nereikia.\u201C "
        "Paaiškinkite, kodėl jis klysta, remdamiesi dviem argumentais: "
        "(a) kodėl vien slaptažodžio nepakanka ir (b) kaip 2FA apsaugo net tada, "
        "kai slaptažodis yra pavogtas. Atsakykite 2\u20133 sakiniais."
    )

    # Q5: Technique 4 (Context Shift: mokykla → biblioteka) — Short answer
    add_body(doc,
        "5. Bibliotekoje naudojatės viešuoju kompiuteriu. Baigę darbą, turite "
        "palikti darbo vietą. Nurodykite tris konkrečius veiksmus, kuriuos "
        "turite atlikti prieš palikdami kompiuterį, ir paaiškinkite, kodėl "
        "kiekvienas jų yra svarbus. Atsakykite sunumeruotu sąrašu.",
        space_after=120
    )

    # ── 3. Internetinės grėsmės ──────────────────────────────────
    add_h3(doc, "3. Internetinės grėsmės")

    # Q6: Technique 2 (Added Complexity) — MCQ
    add_body(doc,
        "6. Gaunate el. laišką su tema \u201ESkubus pranešimas: jūsų paskyra "
        "bus užblokuota per 24 val.\u201C Kuris iš šių požymių yra patikimiausias "
        "phishing atakos rodiklis?"
    )
    add_mcq_option(doc, "A",
        "Laiške yra rašybos klaidų ir neaiškiai suformuluotų sakinių")
    add_mcq_option(doc, "B",
        "Siuntėjo adresas neatitinka oficialaus domeno "
        "(pvz., security@acc0unt-check.xyz)")
    add_mcq_option(doc, "C",
        "Laiškas atsiųstas ne įprastu darbo metu, pavyzdžiui, 2 val. nakties")
    add_mcq_option(doc, "D",
        "Laiške prašoma paspausti nuorodą ir \u201Epatvirtinti savo duomenis\u201C")

    # Q7: Technique 3 (Composite: atpažinimas + reagavimas + pasekmės) — Short answer
    add_body(doc,
        "7. Socialiniame tinkle gaunate žinutę nuo nepažįstamo žmogaus: "
        "\u201ESveiki! Laimėjote 500 EUR dovanų kuponą. Norėdami atsiimti, "
        "spauskite čia ir įveskite savo duomenis.\u201C "
        "Atlikite du dalykus: (a) nurodykite bent du požymius, kodėl ši "
        "žinutė gali būti apgaulė, ir (b) aprašykite, kaip elgtumėtės "
        "pagal trijų žingsnių algoritmą (sustoti, patikrinti, pranešti). "
        "Atsakykite 4\u20135 sakiniais.",
        space_after=120
    )

    # ── 4. Skaitmeninių technologijų poveikis aplinkai ──────────
    add_h3(doc, "4. Skaitmeninių technologijų poveikis aplinkai")

    # Q8: Technique 4 (Context Shift) — MCQ
    add_body(doc,
        "8. Kuris kasdieninis veiksmas sunaudoja daugiausiai energijos "
        "duomenų centruose?"
    )
    add_mcq_option(doc, "A",
        "Vienos trumpos tekstinės žinutės išsiuntimas draugui")
    add_mcq_option(doc, "B",
        "Vieno paveikslėlio parsisiuntimas iš el. pašto priedo")
    add_mcq_option(doc, "C",
        "Valandos trukmės vaizdo įrašo žiūrėjimas aukšta raiška")
    add_mcq_option(doc, "D",
        "Pusvalandžio naršymas tekstiniame tinklalapyje be paveikslėlių")

    # Q9: Technique 1 (Bloom's +1, Apply→Analyze) — Short answer
    add_body(doc,
        "9. Du mokiniai nori sumažinti savo skaitmeninį pėdsaką. Pirmas "
        "planuoja ištrinti senus el. laiškus ir išjungti kompiuterį nakčiai. "
        "Antras planuoja žiūrėti vaizdo įrašus žemesne raiška ir naudoti "
        "mažiau debesijos saugyklų kopijų. Palyginkite šiuos du planus: "
        "kuris labiau sumažins energijos suvartojimą ir kodėl? "
        "Atsakykite 2\u20133 sakiniais."
    )

    # Q10: Technique 2 (Added Complexity) — Short answer
    add_body(doc,
        "10. Paaiškinkite, kodėl netinkamai išmestos e-atliekos kenkia "
        "aplinkai, ir pateikite vieną konkretų pavyzdį, kaip skaitmeninės "
        "technologijos padeda spręsti aplinkosaugos problemas. "
        "Atsakykite 3\u20134 sakiniais.",
        space_after=120
    )

    # ── PATIKRINKITE SAVE ────────────────────────────────────────
    add_h2(doc, "Patikrinkite save")

    checklist = [
        "Ar galiu užtikrintai paaiškinti, kaip taisyklingai įrengti darbo vietą prie kompiuterio?",
        "Ar galiu užtikrintai apibūdinti 20-20-20 taisyklę ir paaiškinti, kodėl ji veikia?",
        "Ar galiu užtikrintai nurodyti bent tris stipraus slaptažodžio požymius?",
        "Ar galiu užtikrintai paaiškinti, kaip 2FA apsaugo paskyrą net pavogus slaptažodį?",
        "Ar galiu užtikrintai atpažinti phishing žinutės požymius?",
        "Ar galiu užtikrintai pritaikyti saugaus reagavimo algoritmą (sustoti, patikrinti, pranešti)?",
        "Ar galiu užtikrintai paaiškinti, kaip skaitmeniniai veiksmai veikia aplinką?",
        "Ar galiu užtikrintai pasiūlyti konkrečius veiksmus savo skaitmeniniam pėdsakui mažinti?",
    ]
    for item in checklist:
        add_checklist_item(doc, item)

    # ── KĄ DARYTI, JEI SUNKU ────────────────────────────────────
    add_h2(doc, "Ką daryti, jei sunku")

    table = doc.add_table(rows=5, cols=2)
    set_table_borders(table)

    # Header row
    hdr_cells = table.rows[0].cells
    write_cell(hdr_cells[0], "Tema", bold=True, color=WHITE, size=11)
    set_cell_shading(hdr_cells[0], "1F4E79")
    write_cell(hdr_cells[1], "Kur peržiūrėti", bold=True, color=WHITE, size=11)
    set_cell_shading(hdr_cells[1], "1F4E79")

    data = [
        ("Ergonomika",
         "001_L, Theory_Pack: sėdėjimo laikysena, ekrano padėtis, 20-20-20 taisyklė, pertraukos"),
        ("Privatumas ir paskyrų sauga",
         "002_L, Theory_Pack: stiprus slaptažodis, 2FA, saugus elgesys su duomenimis"),
        ("Internetinės grėsmės",
         "003_L, Theory_Pack: phishing požymiai, socialinė inžinerija, "
         "sustoti\u2013patikrinti\u2013pranešti"),
        ("Skaitmeninių technologijų poveikis aplinkai",
         "004_L, Theory_Pack: energijos suvartojimas, e-atliekos, skaitmeninis pėdsakas"),
    ]
    for i, (tema, kur) in enumerate(data, 1):
        write_cell(table.rows[i].cells[0], tema, size=11)
        write_cell(table.rows[i].cells[1], kur, size=11)

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Twips(2800)
        row.cells[1].width = Twips(6226)

    # Save
    docx_path = os.path.join(OUTPUT_DIR, "Practice_Task.docx")
    doc.save(docx_path)
    print(f"Saved: {docx_path}")
    return docx_path


if __name__ == "__main__":
    docx_path = build()

    # Convert to PDF
    from docx2pdf import convert
    pdf_path = docx_path.replace(".docx", ".pdf")
    convert(docx_path, pdf_path)
    print(f"Converted: {pdf_path}")

    # Verify
    if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
        print(f"OK: {os.path.getsize(pdf_path)} bytes")
        # Clean up intermediate docx
        os.remove(docx_path)
        print(f"Cleaned up: {docx_path}")
    else:
        print("ERROR: PDF not created or empty")
